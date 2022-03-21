import sys
import logging
import  os
import logging.config
from PyQt5.QtWidgets import QApplication,QWidget,QMainWindow,QAction,QTextBrowser,QMenuBar,QFileSystemModel,QGridLayout,QFileDialog,QVBoxLayout,QSpacerItem
from PyQt5.QtWidgets import QLabel,QHBoxLayout,QPushButton,QCheckBox,QSpinBox,QSizePolicy,QMessageBox
from docx import Document
from docx.opc import exceptions
from TextBrowser import TextBrowser
from ErrorDialog import ErrorDialog
from DocProcess import DocProcess
from pathlib import Path
import lucene
logger = logging.getLogger('logger')


class AppDemo(QWidget):
    def __init__(self):
        super(AppDemo, self).__init__()

        logger.info("initilize application")
        self.setWindowTitle("SmartAss")

        self.fileSystem=QFileSystemModel()
        self.fileSystem.setRootPath("/")


        #工具栏
        self.menuBar=QMenuBar(self)
        #打开文件相关
        fileMenu=self.menuBar.addMenu('文件')
        #保存文件相关
        saveMenu=self.menuBar.addMenu('保存')

        #退出动作
        exit_action=QAction('退出',self)
        exit_action.setShortcut('Ctrl+Q')
        exit_action.triggered.connect(lambda:QApplication.quit())

        #选择要处理的文件
        select_action=QAction('打开...',self)
        select_action.setShortcut('O')
        select_action.triggered.connect(lambda:self.OpenFile())

        #选择停用词列表
        useStopwords_action=QAction('选择停用词列表',self)
        useStopwords_action.setShortcut('Ctrl+S')
        useStopwords_action.triggered.connect(lambda :self.ChooseStopWords())

        fileMenu.addAction(select_action)
        fileMenu.addAction(exit_action)
        fileMenu.addAction(useStopwords_action)

        #保存文件的动作
        saveMainText_action=QAction('保存文档主体信息',self)
        saveMainText_action.triggered.connect(lambda: self.SaveMainText())

        saveSectionText_action = QAction('保存文档章节主体信息', self)
        saveSectionText_action.triggered.connect(lambda: self.SaveSectionText())

        saveTable_action=QAction('保存表格',self)
        saveTable_action.triggered.connect(lambda :self.SaveTable())

        saveMenu.addAction(saveMainText_action)
        saveMenu.addAction(saveSectionText_action)
        saveMenu.addAction(saveTable_action)


        #Maintextbrowser
        textBrowser=TextBrowser(self)
        self.MainTextBrowser=textBrowser
        self.MainTextBrowser.setGeometry(0,25,300,995)
        self.MainTextBrowser.HasTextSignal.connect(self.MainTextBrowserDealer)

        #labels
        keywordsLabel=QLabel(self)
        fileNameLabel=QLabel(self)
        numSectionsLabel=QLabel(self)

        keywordsLabel.setText("关键词:")
        fileNameLabel.setText("文档名:")
        numSectionsLabel.setText("章节数:")

        self.keywordsLabel=keywordsLabel
        self.fileNameLabel=fileNameLabel
        self.numSectionsLabel=numSectionsLabel

        #与上面labels对应的text browser

        self.keywordsBrowser=QTextBrowser(self)
        self.fileNameBrowser=QTextBrowser(self)
        self.numSectionsBrower=QTextBrowser(self)

        #底部的一些部件
        #提取按钮
        self.ExractBtn=QPushButton()
        self.ExractBtn.setText("提取")
        self.ExractBtn.setMaximumSize(200,50)
        self.ExractBtn.clicked.connect(self.Extract)

        #是否用停用词
        self.UseStopwordCheckBox=QCheckBox()
        self.UseStopwordCheckBox.setText("是否使用停用词")
        self.UseStopwordCheckBox.stateChanged.connect(self.StopwordCheckBoxChange)


        #关键词个数设置
        self.NumKeywordsBox=QSpinBox()
        self.NumKeywordsBoxLabel=QLabel()
        self.NumKeywordsBoxLabel.setText("关键词个数")

        #底部部件布局
        bottomSecondLayout=QGridLayout()
        bottomSecondLayout.setContentsMargins(10,10,10,10)
        bottomSecondLayout.addWidget(self.UseStopwordCheckBox,0,5,1,3)
        bottomSecondLayout.addWidget(self.NumKeywordsBoxLabel,1,5,1,1)
        bottomSecondLayout.addWidget(self.NumKeywordsBox,1,6,1,2)

        bottomSpacer=QSpacerItem(100,10,QSizePolicy.Fixed)
        bottomSpacer1 = QSpacerItem(100, 10, QSizePolicy.Fixed)

        bottomLayout=QHBoxLayout()
        bottomLayout.addWidget(self.ExractBtn)
        bottomLayout.addSpacerItem(bottomSpacer)
        bottomLayout.addLayout(bottomSecondLayout)
        bottomLayout.addSpacerItem(bottomSpacer1)

        #布局
        ThirdLayout = QGridLayout()
        ThirdLayout.addWidget(self.keywordsLabel,0,0,1,1)
        ThirdLayout.addWidget(self.fileNameLabel,1,0,1,1)
        ThirdLayout.addWidget(self.numSectionsLabel,2,0,1,1)
        ThirdLayout.addWidget(self.keywordsBrowser,0,1,1,1)
        ThirdLayout.addWidget(self.fileNameBrowser,1,1,1,1)
        ThirdLayout.addWidget(self.numSectionsBrower,2,1,1,1)

        SecondLayout = QGridLayout()
        SecondLayout.addWidget(self.MainTextBrowser,0,0,1,3)
        SecondLayout.addLayout(ThirdLayout,0,4,1,2)

        MainLayout=QVBoxLayout()
        MainLayout.addWidget(self.menuBar)
        MainLayout.addLayout(SecondLayout)
        MainLayout.addLayout(bottomLayout)

        self.setLayout(MainLayout)

        # 待处理的文件
        self.FileOpened=False
        self.FilePath=""
        self.documentInfo=[]
        self.StopWordsPath=""
        self.DocProcess=None
        self.HasTable=False


    def OpenFile(self):
        """
        打开文件
        """
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getOpenFileName(self, "打开word文件", "",
                                                "Docx Files (*.docx)", options=options)
        if(fileName==""):
            return
        text = ""
        logger.info("open file %s",fileName)
        try:#这个是文本文件的情况
            with open(str(fileName), "r") as f:
                text=f.read()
                self.FileOpened=True
                self.FilePath=fileName
        except UnicodeDecodeError:  # 目前只能处理docx文件
                try:
                    document = Document(str(fileName))
                    for para in document.paragraphs:
                        text += para.text + "\n"
                    self.FileOpened=True
                    self.FilePath=fileName
                except (exceptions.PackageNotFoundError):
                    logger.warning("File format incorrect or not exist")
                    dia=ErrorDialog("文件格式不正确或者文件不存在！")
                    dia.exec_()
                    self.FileOpened=False
                    return
        self.MainTextBrowser.clear()
        self.MainTextBrowser.append(text)
        self.fileNameBrowser.clear()
        self.keywordsBrowser.clear()
        self.numSectionsBrower.clear()
        self.HasTable = False


    def StopwordCheckBoxChange(self):
        """
        停用词列表改变后的的唔做
        :return:
        """
        if(self.UseStopwordCheckBox.isChecked()):
            logger.info("CheckBox status: Checked")
            #选取停用词列表
        else:
            logger.info("CheckBox status: UnChecked")
    def ChooseStopWords(self):
        """
        选取停用词列表
        :return:
        """
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getOpenFileName(self, "选取停用词列表", "",
                                                  "TXT Files (*.txt)", options=options)

        try:#这个是文本文件的情况
            with open(str(fileName), "r") as f:
                self.StopWordsPath=fileName
        except (UnicodeDecodeError):  # 目前只能处理docx文件
            logger.warning("File format incorrect or not exist")
            dia = ErrorDialog("文件格式不正确或者文件不存在！")
            dia.exec_()
        except FileNotFoundError:
            return

    def Extract(self):
        """
        处理文件
        :return:
        """
        if(not self.FileOpened):
            logger.info("Extract before open any file")
            dia = ErrorDialog("还未打开文件！")
            dia.exec_()
            return
        if(self.UseStopwordCheckBox.isChecked() and self.StopWordsPath==""):
            dia = ErrorDialog("未选择停用词列表，请选择停用词列表或者取消勾选使用停用词")
            dia.exec_()
            return
        logger.info("begin extraction...")
        if(self.UseStopwordCheckBox.isChecked()):#有停用词列表
            self.DocProcess=DocProcess.DocxProcess(self.FilePath,numwords=int(self.NumKeywordsBox.text()),use_stopwords=True, stopWordsFilePath=self.StopWordsPath)
            self.documentInfo,hasTable= self.DocProcess.ReadDocx()
            self.keywordsBrowser.setText(self.documentInfo[3])
        else:#无停用词列表
            self.DocProcess = DocProcess.DocxProcess(self.FilePath,numwords=int(self.NumKeywordsBox.text()), use_stopwords=False)
            self.documentInfo,hasTable =  self.DocProcess.ReadDocx()
            self.keywordsBrowser.setText(self.documentInfo[3])

        #显示章节数:
        numSections=len(self.documentInfo[5])
        self.numSectionsBrower.setText(str(numSections))

        #显示文档名称
        self.fileNameBrowser.setText(self.documentInfo[0])

        #如果有表格
        if(hasTable):
            self.HasTable=True
            msg = QMessageBox()
            msg.setText("这个文件中含有表格")
            msg.setInformativeText("如果要保存表格请手动保存")
            msg.exec_()


    def MainTextBrowserDealer(self,string):
        """
        处理MainTetxBrowser发来的信号
        :return:
        """
        logger.info("Get file name from MainTextBrowser signal")
        self.FileOpened=True
        self.FilePath=string

    def SaveMainText(self):
        """
        保存文档主体的 csv文件
        :return:
        """
        if(self.DocProcess==None):
            logger.info("Save file info before extract")
            dia = ErrorDialog("您还没有处理文件，无法保存！")
            dia.exec_()
            return

        file_filter = 'csv File ( *.csv )'
        filedia = QFileDialog(self, caption="保存文档主体文件", filter=file_filter, )
        filedia.setLabelText(QFileDialog.Accept,"Save")
        filedia.fileSelected.connect(self.SaveMain)
        filedia.exec_()




    def SaveSectionText(self):
        """
        保存章节主体的 csv文件
        :return:
        """
        if (self.DocProcess == None):
            logger.info("Save file info before extract")
            dia = ErrorDialog("您还没有处理文件，无法保存！")
            dia.exec_()
            return
        file_filter = 'csv File ( *.csv )'
        filedia=QFileDialog(self, caption="保存章节主体文件", filter=file_filter, )
        filedia.setLabelText(QFileDialog.Accept, "Save")
        filedia.fileSelected.connect(self.SaveSection)
        filedia.exec_()

    def SaveSection(self,s):
        """
        保存文档章节主体信息
        :param s: 文档路径
        :return:
        """
        if(not Path(s).exists()):
            #不存在直接保存
            logger.info("Saving file info to %s", s)
            self.DocProcess.SaveCsvSection(s, self.documentInfo, Path(s).exists())
        else:
            #先询问是否要覆盖
            msg=QMessageBox()
            msg.setText("文件已经存在")
            msg.setInformativeText("您可以选择追加或者覆盖")
            overWriteBtn=msg.addButton("覆盖",QMessageBox.ActionRole)
            appendBtn=msg.addButton("追加",QMessageBox.ActionRole)
            cancleBtn=msg.addButton("取消",QMessageBox.ActionRole)
            msg.exec_()
            if(msg.clickedButton()==overWriteBtn):
                logger.info("Overwrite file: %s",s)
                self.DocProcess.SaveCsvSection(s, self.documentInfo, False)
            if(msg.clickedButton()==appendBtn):
                logger.info("Append file: %s",s)
                self.DocProcess.SaveCsvSection(s, self.documentInfo, True)
            if(msg.clickedButton()==cancleBtn):
                # 这里目前只能return 会导致filedialog也关闭，目前没有找到解决方法
                return

    def SaveMain(self,s):
        """
        保存文档主体信息
        :param s: 文档路径
        :return:
        """
        if(not Path(s).exists()):
            #不存在直接保存
            logger.info("Saving file info to %s", s)
            self.DocProcess.SaveCsv(s, self.documentInfo, Path(s).exists())
        else:
            #先询问是否要覆盖
            msg=QMessageBox()
            msg.setText("文件已经存在")
            msg.setInformativeText("您可以选择追加或者覆盖")
            overWriteBtn=msg.addButton("覆盖",QMessageBox.ActionRole)
            appendBtn=msg.addButton("追加",QMessageBox.ActionRole)
            cancleBtn=msg.addButton("取消",QMessageBox.ActionRole)
            msg.exec_()
            if(msg.clickedButton()==overWriteBtn):
                logger.info("Overwrite file: %s",s)
                self.DocProcess.SaveCsv(s, self.documentInfo, False)
            if(msg.clickedButton()==appendBtn):
                logger.info("Append file: %s",s)
                self.DocProcess.SaveCsv(s, self.documentInfo, True)
            if(msg.clickedButton()==cancleBtn):
                # 这里目前只能return 会导致filedialog也关闭，目前没有找到解决方法
                return

    def SaveTable(self):
        if (self.DocProcess == None):
            logger.info("Save file info before extract")
            dia = ErrorDialog("您还没有处理文件，无法保存！")
            dia.exec_()
            return
        if (self.HasTable == False):
            logger.info("No table can be saved")
            dia = ErrorDialog("这个文件中不含表格！")
            dia.exec_()
            return

        file_filter = 'csv File ( *.csv );;xlsx Files (*.xlsx)'
        filedia = QFileDialog(self, caption="保存表格", filter=file_filter, )
        filedia.setLabelText(QFileDialog.Accept, "Save")
        filedia.fileSelected.connect(self._SaveTable)
        filedia.exec_()

    def _SaveTable(self,s):
        if (not Path(s).exists()):
            # 不存在直接保存
            logger.info("Saving file info to %s", s)
            for table in self.DocProcess.tables:
                self.DocProcess.SaveTable(s, table, False)
        else:
            if (os.path.splitext(s)[-1] == ".csv"):
                # csv文件可以覆盖或者追加
                # 先询问是否要覆盖
                msg = QMessageBox()
                msg.setText("文件已经存在")
                msg.setInformativeText("您可以选择追加或者覆盖")
                overWriteBtn = msg.addButton("覆盖", QMessageBox.ActionRole)
                appendBtn = msg.addButton("追加", QMessageBox.ActionRole)
                cancleBtn = msg.addButton("取消", QMessageBox.ActionRole)
                msg.exec_()
                if (msg.clickedButton() == overWriteBtn):
                    logger.info("Overwrite table file: %s", s)
                    for table in self.DocProcess.tables:
                        self.DocProcess.SaveTable(s, table, False)
                    return
                if (msg.clickedButton() == appendBtn):
                    logger.info("Append table file: %s", s)
                    for table in self.DocProcess.tables:
                        self.DocProcess.SaveTable(s, table, True)
                    return
                if (msg.clickedButton() == cancleBtn):
                    # 这里目前只能return 会导致filedialog也关闭，目前没有找到解决方法
                    return
            elif (os.path.splitext(s)[-1] == ".xlsx"):
                # xlsx文件只能覆盖
                msg = QMessageBox()
                msg.setText("文件已经存在")
                msg.setInformativeText("您确定要覆盖此文件吗")
                overWriteBtn = msg.addButton("覆盖", QMessageBox.ActionRole)
                cancleBtn = msg.addButton("取消", QMessageBox.ActionRole)
                msg.exec_()
                if (msg.clickedButton() == overWriteBtn):
                    logger.info("Overwrite table file: %s", s)
                    for table in self.DocProcess.tables:
                        self.DocProcess.SaveTable(s, table, False)
                    return
                if (msg.clickedButton() == cancleBtn):
                    # 这里目前只能return 会导致filedialog也关闭，目前没有找到解决方法
                    return

            logger.error("file format worng!")



if __name__ == '__main__':
    #初始化日志
    logging.config.fileConfig('./log.conf')

    app=QApplication(sys.argv)
    demo=AppDemo()
    demo.resize(800,600)
    demo.show()
    try:
        sys.exit(app.exec_())
    except SystemExit:
        logger.info("close window")