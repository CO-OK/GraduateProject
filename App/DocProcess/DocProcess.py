from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv
# from . import TextRank
from Code.App.TextRank import TextRank
import os
import xlwt
from Code.App.DocProcess.LexicalChains import Summarizer,LexicalChain
class DocxProcess:
    def __init__(self,docxFilePath,numwords=10, windows=2,
                 use_stopwords=False, stopWordsFilePath=None,
                 pr_config={'alpha': 0.85, 'max_iter': 100}):
        """
        :param docxFilePath docx文档路径
        :param numwords: 关键词个数
        :param windows: 提取关键词时邻接矩阵的滑动窗口数
        :param use_stopwords: 是否使用停用词
        :param stopWordsFilePath:停用词列表的路径
        :param pr_config:提取关键词时的配置
        """
        self.DocxFilePath=docxFilePath
        self.NumWords=numwords
        self.windows=windows
        self.UseStopwords=use_stopwords
        self.StopwordsPath=stopWordsFilePath
        self.PrConfig=pr_config

        #表格
        self.tables=None

    def GetKeyWords(self,text):
        """
        :param text:要提取关键词的文本2962466400
        :return: 关键词字符串
        """
        textRank=TextRank.TextRank(text,self.windows,self.UseStopwords,self.StopwordsPath,False,self.PrConfig)
        keywords=textRank.get_n_keywords(self.NumWords)
        res=""
        for item in keywords:
            res+=item[0]+" "
        return res


    def ReadDocx(self):
        """
        :param filePath: docx文件路径
        :return: 标题、文档编号、文档正文、文档关键词、正文文件(路径)、各个章节组成的列表
        """
        document = Document(self.DocxFilePath)
        Title=""#标题
        Doucment_num=""#文档编号
        Text = ""#文档正文
        #获取标题
        Document_num_index=0
        #从文章开始能够找到的所有居中的段落都算标题
        for i,paragraph in enumerate(document.paragraphs):
            if(paragraph.alignment==WD_ALIGN_PARAGRAPH.CENTER):
                Title+=paragraph.text+" "
                Document_num_index=i
            else:
                break
        #如果标题还为空则自动采用文件名作为标题
        if(Title==""):
            Title=os.path.basename(self.DocxFilePath)
        #获取文档编号
        #按照安全总管系列的文章来说标题的最后一部分应该就是文章编号
        Doucment_num=document.paragraphs[Document_num_index].text

        #文档正文
        for paragraph in document.paragraphs:
            Text+=paragraph.text+"\n"

        #获取文档章节
        Sections=[]
        section=[]
        for i,paragraph in enumerate(document.paragraphs):
            if(i<=Document_num_index):
                continue
            else:#进入正文
                if(paragraph.text==""):continue#这段没有文字的话就继续
                if(paragraph.runs[0].bold):#加粗是章节标题
                    Sections.append(section)
                    section=[]
                    section.append(paragraph.text)
                else:
                    section.append(paragraph.text)
        if(len(section)!=0):
            Sections.append(section)
        final=[]
        final.append(Title)
        final.append(Doucment_num)
        final.append(Text)
        keywords=self.GetKeyWords(Text)#关键词
        final.append(keywords)
        final.append(self.DocxFilePath)
        final.append(Sections)
        # 处理表格
        if(len(document.tables)!=0):#有表格
            self.tables=document.tables
            return final,True

        return final,False


    def SaveCsv(self,path,documentInfo,exist=False):
        """
        :param path:最后csv文件的保存路径
        :param documentInfo:已经解析好的文档相关信息
        :return:
        """

        if(exist):
            arg='a+'
        else:
            arg='w'
        with open(path,arg) as f:
            #后缀必须是csv才可以被创建
            csv_write = csv.writer(f)
            #先写入标题
            if(not exist):
                csv_head = ["文档编号","文档名称","正文","正文关键词","文档正文文件","文档章节数","文档附件","文档附件关键词","文档附件文件"]
                csv_write.writerow(csv_head)
            #再写入内容
            row=[]
            row.append(documentInfo[1])
            row.append(documentInfo[0])
            row.append(documentInfo[2])
            row.append(documentInfo[3])
            row.append(documentInfo[4])
            #关键词
            # textRank = TextRank.TextRank(str(documentInfo[2]), 2, True, "./TextRank/cn_stopwords.txt", False)
            # keyWords=textRank.get_n_keywords(10)#list of tuple
            #变为字符串
            # keyWordsList=[item[0] for item in keyWords]
            # row.append(" ".join(keyWordsList))
            # row.append("")
            row.append(len(documentInfo[5]))
            row.append("")
            row.append("")
            row.append("")
            csv_write.writerow(row)

    def SaveCsvSection(self,path,documentInfo,exist=False):
        # 文档篇章结构信息保存
        if (exist):
            arg = 'a+'
        else:
            arg = 'w'
        with open(path,arg) as f:
            #后缀必须是csv才可以被创建
            csv_write = csv.writer(f)
            if(not exist):
                csv_head = ["文档编号", "文档名称", "文档章节题目","文档章节摘要","文档章节关键词","文档章节号","文档章节内容"]
                csv_write.writerow(csv_head)
            # 按文档章节来进行遍历
            Sections = documentInfo[5]
            for i,section in enumerate(Sections):
                row=[]
                # 文档编号
                row.append(documentInfo[1])
                # wendangmingcheng
                row.append(documentInfo[0])
                # 文档章节题目 section的第一个
                row.append(section[0])
                # 摘要暂时为空
                # row.append("")
                # 章节关键词
                text=""
                for sentence in section:
                    text+=sentence
                keywords=self.GetKeyWords(text)
                # 通过text获取摘要
                chain=LexicalChain(text)
                fs =Summarizer()
                final_chain=chain.get_final_chain()
                summary=""
                if len(chain.sentence) >= 5:
                    n = 5
                else:
                    n = 2
                for s in fs.summarize(chain.sentence, final_chain, n):
                    print(s)
                    summary+=s+'\n'
                row.append(summary)
                row.append(keywords)
                #章节号
                row.append(str(i))
                #内容
                row.append(text)
                csv_write.writerow(row)

    def SaveTable(self,path, table, exist=False):
        """
        保存表格
        :param path:保存的文件路径
        :param table:表格object
        :param exist:是否追加
        :return:
        """
        if (exist):
            arg = 'a+'
        else:
            arg = 'w'
        # 根据后缀来选择存储的文件
        if (os.path.splitext(path)[-1] == ".csv"):  # csv形式
            rows = []
            for row in table.rows:
                row_element = []
                for cell in row.cells:
                    row_element.append(cell.text)
                rows.append(row_element)
            with open(path, arg) as f:
                csv_write = csv.writer(f)
                csv_write.writerows(rows)
        elif (os.path.splitext(path)[-1] == ".xlsx"):  # excel形式
            workbook = xlwt.Workbook(encoding='utf-8')
            worksheet = workbook.add_sheet("table", cell_overwrite_ok=True)
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    worksheet.write(i, j, cell.text)
            # 合并单元格 这部分耗时太长，要改进
            for col_num, col in enumerate(table.columns):
                row_num_start = 0
                row_num_end = row_num_start + 1
                # print(col_num)
                while (row_num_start < len(col.cells)):
                    text = table.cell(row_num_start, col_num).text
                    while (row_num_end < len(col.cells) and text == table.cell(row_num_end, col_num).text):
                        row_num_end += 1
                    if (row_num_start + 1 < row_num_end):
                        worksheet.write_merge(row_num_start, row_num_end - 1, col_num, col_num, text)
                        row_num_start = row_num_end
                        row_num_end += 1
                    else:
                        row_num_start += 1
                        row_num_end += 1

            workbook.save(path)





# if __name__ == "__main__":
#     document = Document("../../Data/test.docx")
#
#     SaveTable("../../Data/ttt.xlsx",document.tables[0])
#
#     # doc= DocxProcess("../../Data/test.docx")
#     # doc.ReadDocx()
#     pass



